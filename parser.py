from docx import Document
import json
import os

planner = {}
directory = os.fsencode('./planners')
period = int(input(f'What period:\n'))
for file in os.listdir(directory):
    name = os.fsdecode(file)
    subject = name.split('.')[0]
    starting_row = 5
    table_index = 0

    if name.endswith('.docx'):
        document = Document(os.path.join('./planners', name))

        for i in range(0, len(document.tables)):
            for j in range(0, len(document.tables[i].rows)):
                row = document.tables[i].rows[j]
                if row.cells[0].text == 'Periode':
                    if row.cells[2].text == f'{period}' or row.cells[2].text == f'{4 + period}':
                        table_index = i

        for i in range(0, len(document.tables[table_index].rows)):
            if document.tables[table_index].rows[i].cells[4].text == 'Wat moet je afhebben/leren':
                starting_row = i + 1

        for i in range(starting_row, len(document.tables[table_index].rows)):
            row = document.tables[table_index].rows[i]
            date = row.cells[1].text.replace(' ', '').replace('\n', '')
            todo_home = row.cells[4].text.replace('\n', ' ')
            todo_class = row.cells[3].text.replace('\n', ' ')

            if date and date in planner:
                planner[date].append({ 'todo_home': todo_home, 'todoClass': todo_class, 'subject': subject })
            elif date and date not in planner:
                planner[date] = [{'todo_home': todo_home, 'todo_class': todo_class, 'subject': subject }]

with open('data.json', 'w') as data:
    data.seek(0)
    json.dump(planner, data, indent=4)
    data.truncate()

